from typing import Optional, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
import sqlite3
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.prompts.chat import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import AIMessage, HumanMessage, SystemMessage
from fastapi import FastAPI, HTTPException
import uvicorn
import boto3
from fnmatch import fnmatchcase
import json
from typing import List, Tuple
import os
from docx.oxml.ns import qn
from docx import Document as DocxDocument
from io import BytesIO

os.environ['OPENAI_API_KEY'] = 'my-api-key'

model_name = "gpt-4o"
temperature = 0
llm = ChatOpenAI(model=model_name, temperature=temperature)

embeddings = OpenAIEmbeddings()

current_user = 'TECHNONIKOL'

# Настройка клиента для Yandex S3
session = boto3.session.Session()
s3_client = session.client(
    service_name='s3',
    endpoint_url='https://storage.yandexcloud.net',
    aws_access_key_id='my_access_key_id',
    aws_secret_access_key='my_secret_access_key',
)

CHROMA_PATH = f'./chroma/{current_user}/new/'


def init_metadata_db():
    with sqlite3.connect('metadata.db') as conn:
        conn.execute('''
        CREATE TABLE IF NOT EXISTS uploaded_docs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        global_source TEXT,
        filename TEXT
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS history_messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT,
                user_type TEXT,
                message TEXT,
                tmstmp DATETIME DEFAULT CURRENT_TIMESTAMP
        );
        ''')


init_metadata_db()


class SQLiteChatHistory():
    def __init__(self, db_path="metadata.db"):
        self.db_path = db_path

    def add_message(self, message):
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        if isinstance(message, HumanMessage):
            user_type = "human"
            message = message.content
        elif isinstance(message, AIMessage):
            user_type = "ai"
            message = message.content
        elif isinstance(message, SystemMessage):
            user_type = "system"
            message = message.content
        else:
            raise ValueError("Invalid message type")
        c.execute("INSERT INTO history_messages (user_id, user_type, message) VALUES (?, ?, ?)",
                  (current_user, user_type, message))
        conn.commit()
        conn.close()

    def messages(self, limit=15):
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(f"SELECT * FROM history_messages WHERE user_id = '{current_user}' ORDER BY id DESC LIMIT {limit}")
        resp = c.fetchall()[::-1]
        chat_history = []
        for row in resp:
            id, user_id, user_type, message, tmstmp = row
            if user_type == "human":
                chat_history.append(HumanMessage(content=message))
            elif user_type == "ai":
                chat_history.append(AIMessage(content=message))
            elif user_type == "system":
                chat_history.append(SystemMessage(content=message))
        conn.commit()
        conn.close()
        messages = ChatMessageHistory(messages=chat_history)
        return messages

    def delete_chat_history_last_n(self, n=10):
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(f'''
        with max_id as (select max(id) as maxid from history_messages where user_id = '{current_user}')
        DELETE FROM history_messages
        WHERE id BETWEEN (select maxid from max_id) - {n} AND (select maxid from max_id)
        ''')
        conn.commit()
        conn.close()


def add_filename_to_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''INSERT INTO uploaded_docs (global_source, filename) values ('{source}', '{filename}') ; ''')


def delete_filename_from_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''DELETE from uploaded_docs where global_source = '{source}' and filename ='{filename}' ; ''')


def get_uploaded_filenames(global_source):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"select * from uploaded_docs where global_source = '{global_source}' ")
        rows = cursor.fetchall()
    unique_list = [f"{row[2]}" for row in rows]
    return unique_list


class Document:
    def __init__(self, source: str, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.source = source
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {'source': source}


def get_uploaded_filenames(source) -> List[str]:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT filename FROM uploaded_docs WHERE global_source = ?", (source,))
        rows = cursor.fetchall()
    filenames = [row[0] for row in rows]
    return filenames


def load_s3_files(bucket: str, prefix: str, suffix: str) -> List[str]:
    """List files in a given S3 bucket with a specified prefix and suffix."""
    try:
        response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
        files = [content['Key'] for content in response.get('Contents', []) if content['Key'].endswith(suffix)]
        if not files:
            print(f"No files found in bucket {bucket} with prefix {prefix} and suffix {suffix}")
        else:
            print(f"Files found in bucket {bucket} with prefix {prefix} and suffix {suffix}: {files}")
        return files
    except Exception as e:
        print(f"Error listing files in bucket {bucket} with prefix {prefix} and suffix {suffix}: {e}")
        return []


def load_docx_new(source, bucket: str) -> List[Document]:
    prefix = 'TECHNONIKOL/docx/'
    suffix = '.docx'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read()

                # Используем BytesIO для чтения содержимого файла как бинарного потока
                doc_stream = BytesIO(content)
                doc = DocxDocument(doc_stream)

                # Извлекаем текст из документа docx
                full_text = []
                image_counter = 1

                # Получаем имя файла без расширения и создаем соответствующую папку
                filename_without_extension = os.path.splitext(os.path.basename(file))[0]
                image_folder = filename_without_extension  # Используем оригинальное имя файла для папки

                for para in doc.paragraphs:
                    # Обработка параграфов для создания ссылок на изображения
                    para_text = para.text
                    for run in para.runs:
                        for drawing in run.element.findall('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            image_rId = drawing.get(qn('r:embed'))
                            image_part = doc.part.related_parts[image_rId]
                            image_filename = f'image_{image_counter:02d}.{image_part.content_type.split("/")[-1]}'
                            image_counter += 1

                            # Загрузка изображения в бакет Яндекса
                            img_content = image_part.blob
                            s3_image_key = f"A100/images/{image_folder}/{image_filename}"
                            s3_client.put_object(
                                Bucket=bucket,
                                Key=s3_image_key,
                                Body=img_content,
                                ContentDisposition='inline',
                                ContentType=image_part.content_type
                            )

                            # Генерация URL для изображения
                            s3_image_url = f"https://storage.yandexcloud.net/{bucket}/{s3_image_key}"
                            para_text += f'\n{s3_image_url}'
                    full_text.append(para_text)
                content = '\n'.join(full_text)

                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading docx file {file}: {e}")

    return docs if docs else None


def load_txts(source, bucket: str) -> List[Document]:
    prefix = 'TECHNONIKOL/txt/'
    suffix = '.txt'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read().decode('utf-8')
                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading txt file {file}: {e}")

    return docs if docs else None


def load_jsons(source, bucket: str) -> Tuple[List[Document], List[dict]]:
    prefix = 'TECHNONIKOL/json/'
    suffix = '.json'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    json_docs, json_metadata = [], []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = json.loads(obj['Body'].read().decode('utf-8'))
                json_docs.append(content)
                json_metadata.append({'source': file})
            except Exception as e:
                print(f"Error reading json file {file}: {e}")

    return (json_docs, json_metadata) if json_docs else (None, None)


def load_documents(global_source, bucket: str, file_types: List[str]) -> dict:
    """
    Загружаем документы в зависимости от типа документа из Yandex S3
    """
    all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None}
    if 'txt' in file_types:
        txt_docs = load_txts(global_source, bucket)
        all_docs['txt'] = txt_docs
    if 'json' in file_types:
        json_docs, json_metadata = load_jsons(global_source, bucket)
        all_docs['json'] = json_docs
        all_docs['json_metadata'] = json_metadata
    if 'docx' in file_types:
        docx_docs = load_docx_new(global_source, bucket)
        all_docs['docx'] = docx_docs
    return all_docs


# Пример использования
DATA_BUCKET = 'utlik'
DOCS = load_documents('local', DATA_BUCKET, ['txt', 'json', 'docx'])


def split_docs_to_chunks(documents: dict, file_types: List[str], chunk_size=2000, chunk_overlap=200):
    all_chunks = []
    if 'txt' in file_types and documents['txt'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['txt']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    if 'json' in file_types and documents['json'] is not None:
        json_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        json_chunks = json_splitter.create_documents([json.dumps(doc, ensure_ascii=False) for doc in documents['json']],
                                                     metadatas=documents['json_metadata'])
        all_chunks.extend(json_chunks)

    if 'docx' in file_types and documents['docx'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['docx']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    return all_chunks


chunks_res = split_docs_to_chunks(DOCS, ['txt', 'json', 'docx'])


def get_chroma_vectorstore(documents, embeddings, persist_directory):
    if os.path.isdir(persist_directory) and os.listdir(persist_directory):
        print("Loading existing Chroma vectorstore...")
        vectorstore = Chroma(
            embedding_function=embeddings, persist_directory=persist_directory
        )

        existing_files = get_uploaded_filenames('local')
        uniq_sources_to_add = set(
            doc.metadata['source'] for doc in chunks_res
            if doc.metadata['source'] not in existing_files
        )

        if uniq_sources_to_add:
            vectorstore.add_documents(
                documents=[doc for doc in chunks_res if doc.metadata['source'] in uniq_sources_to_add],
                embedding=embeddings
            )
            for filename in uniq_sources_to_add:
                add_filename_to_metadata('local', filename)
        else:
            print('Новых документов не было, пропускаем шаг добавления')

    else:
        print("Creating and indexing new Chroma vectorstore...")
        vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=embeddings, persist_directory=persist_directory
        )
        uniq_sources_to_add = set(doc.metadata['source'] for doc in documents)
        for filename in uniq_sources_to_add:
            add_filename_to_metadata('local', filename)

    return vectorstore


vectorstore = get_chroma_vectorstore(documents=chunks_res, embeddings=embeddings, persist_directory=CHROMA_PATH)
retriever = vectorstore.as_retriever(search_kwargs={"k": 2}, search_type='similarity')


def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


chat_history_for_chain = SQLiteChatHistory()

prompt_new = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            '''
            Please find the most relevant products for the query from the provided knowledge base {context}.
            If possible, utilize information from previous messages. If uncertain of the answer, briefly indicate so.
            Respond to the following question using only your existing knowledge and the provided knowledge base: {context}
            Avoid using the internet to find information.
            Always ensure three distinct and unique answers are provided for each product query.
            Each product should be distinctly different and not repeat others.
            Include the product name and a URL exclusively sourced from the provided knowledge base.
            Do not create new URLs. Use only the URLs provided in the knowledge base.
            If a URL from the knowledge base is unavailable, acknowledge this and refrain from providing a URL.

            ''',
        ),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}"),
    ]
)

chain_new = prompt_new | llm

chain_with_message_history = RunnableWithMessageHistory(
    chain_new,
    lambda session_id: chat_history_for_chain.messages(limit=10),
    input_messages_key="question",
    history_messages_key="chat_history",
)

app = FastAPI()


@app.post("/rag_chat/")
async def ask_question(question_data: dict = None):
    if question_data is None:
        raise HTTPException(status_code=400, detail="Question data is required")

    question = question_data.get('question', None)
    if question is None:
        raise HTTPException(status_code=400, detail="Question is required")

    answer = chain_with_message_history.invoke(
        {"question": question, "context": format_docs(retriever.invoke(question))},
        {"configurable": {"session_id": 1}})
    answer = answer.content
    if answer is not None:
        chat_history_for_chain.add_message(HumanMessage(content=question))
        chat_history_for_chain.add_message(AIMessage(content=answer))

    return {"answer": answer}


if __name__ == "__main__":
    uvicorn.run("tech:app", host="0.0.0.0", port=8123, reload=True)


